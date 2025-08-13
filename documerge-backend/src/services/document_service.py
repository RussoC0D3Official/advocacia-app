import os
import json
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Inches
from google.cloud import storage
from src.models.user import db, GeneratedPetition, Question, ThesisQuestionLink, Thesis

class DocumentService:
    def __init__(self):
        self.bucket_name = 'documerge-storage'
        self.client = None
        self.bucket = None
        
        # Inicializa cliente GCS se as credenciais estiverem disponíveis
        try:
            self.client = storage.Client(project="documerge-api")
            self.bucket = self.client.bucket(self.bucket_name)
        except Exception as e:
            print(f"Aviso: Google Cloud Storage não configurado: {e}")
            # Em desenvolvimento, usa armazenamento local
            self.local_storage_path = '/tmp/advocacia_documents'
            os.makedirs(self.local_storage_path, exist_ok=True)
    
    def upload_thesis_file(self, file, client_id, title):
        """Faz upload de um arquivo de tese.
        Aceita:
        - file: FileStorage (tem método save) OU
        - file: caminho str para arquivo local OU
        - file: file-like (read())
        """
        try:
            # Gera nome único para o arquivo
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"client_{client_id}/theses/{timestamp}_{title.replace(' ', '_')}.docx"
            
            if self.bucket:
                # Upload para GCS
                blob = self.bucket.blob(filename)
                if hasattr(file, 'read') and not hasattr(file, 'save'):
                    blob.upload_from_file(file, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                elif isinstance(file, str) and os.path.exists(file):
                    with open(file, 'rb') as f:
                        blob.upload_from_file(f, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                else:
                    blob.upload_from_file(file, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                return f"gs://{self.bucket_name}/{filename}"
            else:
                # Armazenamento local para desenvolvimento
                local_path = os.path.join(self.local_storage_path, filename)
                os.makedirs(os.path.dirname(local_path), exist_ok=True)
                if hasattr(file, 'save'):
                    file.save(local_path)
                elif isinstance(file, str) and os.path.exists(file):
                    import shutil
                    shutil.copy2(file, local_path)
                elif hasattr(file, 'read'):
                    with open(local_path, 'wb') as f:
                        f.write(file.read())
                else:
                    raise Exception('Tipo de arquivo não suportado para upload')
                return local_path
                
        except Exception as e:
            raise Exception(f"Erro ao fazer upload do arquivo: {str(e)}")
    
    def upload_petition_file(self, content, user_id, client_id, title):
        """Salva uma petição gerada"""
        try:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"client_{client_id}/petitions/{timestamp}_{title.replace(' ', '_')}.docx"
            
            if self.bucket:
                # Upload para GCS
                blob = self.bucket.blob(filename)
                blob.upload_from_string(content, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                return f"gs://{self.bucket_name}/{filename}"
            else:
                # Armazenamento local
                local_path = os.path.join(self.local_storage_path, filename)
                os.makedirs(os.path.dirname(local_path), exist_ok=True)
                with open(local_path, 'wb') as f:
                    f.write(content)
                return local_path
                
        except Exception as e:
            raise Exception(f"Erro ao salvar petição: {str(e)}")
    
    def download_file(self, gcs_path):
        """Baixa um arquivo do GCS ou local"""
        try:
            if gcs_path.startswith('gs://'):
                # Download do GCS
                if not self.bucket:
                    raise Exception("Google Cloud Storage não configurado")
                
                blob_name = gcs_path.replace(f"gs://{self.bucket_name}/", "")
                blob = self.bucket.blob(blob_name)
                
                # Baixa para arquivo temporário
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
                blob.download_to_filename(temp_file.name)
                return temp_file.name
            else:
                # Arquivo local
                if os.path.exists(gcs_path):
                    return gcs_path
                else:
                    raise Exception(f"Arquivo não encontrado: {gcs_path}")
                    
        except Exception as e:
            raise Exception(f"Erro ao baixar arquivo: {str(e)}")
    
    def delete_file(self, gcs_path):
        """Remove um arquivo"""
        try:
            if gcs_path.startswith('gs://'):
                # Remove do GCS
                if self.bucket:
                    blob_name = gcs_path.replace(f"gs://{self.bucket_name}/", "")
                    blob = self.bucket.blob(blob_name)
                    blob.delete()
            else:
                # Remove arquivo local
                if os.path.exists(gcs_path):
                    os.remove(gcs_path)
                    
        except Exception as e:
            print(f"Aviso: Erro ao remover arquivo {gcs_path}: {e}")
    
    def generate_petition(self, petition_model_id, form_answers, user_id, client_id, title, process_number=None):
        """Gera uma petição baseada nas respostas do formulário (UC-01)"""
        try:
            # Busca as perguntas do modelo
            questions = Question.query.filter_by(petition_model_id=petition_model_id).order_by(Question.order).all()
            
            # Identifica as teses relevantes baseadas nas respostas
            selected_theses = []
            
            for question in questions:
                question_id = str(question.id)
                if question_id in form_answers:
                    answer = 'sim' if form_answers[question_id] else 'nao'
                    
                    # Busca teses vinculadas a esta resposta
                    links = ThesisQuestionLink.query.filter_by(
                        question_id=question.id,
                        answer=answer
                    ).all()
                    
                    for link in links:
                        if link.thesis not in selected_theses:
                            selected_theses.append(link.thesis)
            
            if not selected_theses:
                raise Exception("Nenhuma tese foi selecionada com base nas respostas fornecidas")
            
            # Cria o documento final
            final_doc = Document()
            
            # Adiciona título
            title_paragraph = final_doc.add_heading(title, 0)
            
            # Adiciona número do processo se fornecido
            if process_number:
                final_doc.add_paragraph(f"Processo nº: {process_number}")
            
            final_doc.add_paragraph()  # Linha em branco
            
            # Mescla as teses selecionadas
            for i, thesis in enumerate(selected_theses):
                # Baixa o arquivo da tese
                thesis_file_path = self.download_file(thesis.gcs_path)
                
                try:
                    # Abre o documento da tese
                    thesis_doc = Document(thesis_file_path)
                    
                    # Adiciona cabeçalho da tese
                    final_doc.add_heading(f"{i+1}. {thesis.title}", 1)
                    
                    # Copia o conteúdo da tese
                    for paragraph in thesis_doc.paragraphs:
                        if paragraph.text.strip():  # Ignora parágrafos vazios
                            new_paragraph = final_doc.add_paragraph(paragraph.text)
                            # Tenta preservar formatação básica
                            for run in paragraph.runs:
                                if run.bold:
                                    new_paragraph.runs[-1].bold = True
                                if run.italic:
                                    new_paragraph.runs[-1].italic = True
                    
                    # Adiciona espaço entre teses
                    final_doc.add_paragraph()
                    
                finally:
                    # Remove arquivo temporário se foi baixado do GCS
                    if thesis.gcs_path.startswith('gs://') and os.path.exists(thesis_file_path):
                        os.remove(thesis_file_path)
            
            # Salva o documento final
            temp_output = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            final_doc.save(temp_output.name)
            
            # Lê o conteúdo do arquivo
            with open(temp_output.name, 'rb') as f:
                content = f.read()
            
            # Remove arquivo temporário
            os.remove(temp_output.name)
            
            # Faz upload da petição gerada
            gcs_path = self.upload_petition_file(content, user_id, client_id, title)
            
            # Salva metadados no banco
            petition = GeneratedPetition(
                user_id=user_id,
                client_id=client_id,
                title=title,
                process_number=process_number,
                gcs_path=gcs_path,
                form_data=json.dumps(form_answers)
            )
            
            db.session.add(petition)
            db.session.commit()
            
            return petition
            
        except Exception as e:
            db.session.rollback()
            raise Exception(f"Erro ao gerar petição: {str(e)}")
    
    def get_petition_content(self, petition_id):
        """Retorna o conteúdo de uma petição para visualização/edição (UC-02)"""
        try:
            petition = GeneratedPetition.query.get(petition_id)
            if not petition:
                raise Exception("Petição não encontrada")
            
            # Baixa o arquivo
            file_path = self.download_file(petition.gcs_path)
            
            try:
                # Abre o documento
                doc = Document(file_path)
                
                # Extrai o texto
                content = []
                for paragraph in doc.paragraphs:
                    content.append(paragraph.text)
                
                return {
                    'petition': petition.to_dict(),
                    'content': '\n'.join(content)
                }
                
            finally:
                # Remove arquivo temporário se necessário
                if petition.gcs_path.startswith('gs://') and os.path.exists(file_path):
                    os.remove(file_path)
                    
        except Exception as e:
            raise Exception(f"Erro ao obter conteúdo da petição: {str(e)}")
    
    def update_petition_content(self, petition_id, new_content, new_title=None):
        """Atualiza o conteúdo de uma petição (UC-02)"""
        try:
            petition = GeneratedPetition.query.get(petition_id)
            if not petition:
                raise Exception("Petição não encontrada")
            
            # Cria novo documento com o conteúdo atualizado
            doc = Document()
            
            # Adiciona título
            title = new_title or petition.title
            doc.add_heading(title, 0)
            
            # Adiciona número do processo se existir
            if petition.process_number:
                doc.add_paragraph(f"Processo nº: {petition.process_number}")
            
            doc.add_paragraph()  # Linha em branco
            
            # Adiciona o conteúdo (divide por linhas)
            for line in new_content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
            
            # Salva temporariamente
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_file.name)
            
            # Lê o conteúdo
            with open(temp_file.name, 'rb') as f:
                content = f.read()
            
            # Remove arquivo temporário
            os.remove(temp_file.name)
            
            # Remove arquivo antigo
            self.delete_file(petition.gcs_path)
            
            # Faz upload do novo arquivo
            new_gcs_path = self.upload_petition_file(
                content, petition.user_id, petition.client_id, title
            )
            
            # Atualiza no banco
            petition.gcs_path = new_gcs_path
            if new_title:
                petition.title = new_title
            petition.updated_at = datetime.utcnow()
            
            db.session.commit()
            
            return petition
            
        except Exception as e:
            db.session.rollback()
            raise Exception(f"Erro ao atualizar petição: {str(e)}")
    
    def list_user_petitions(self, user_id):
        """Lista petições de um usuário"""
        try:
            petitions = GeneratedPetition.query.filter_by(user_id=user_id).order_by(GeneratedPetition.created_at.desc()).all()
            return [petition.to_dict() for petition in petitions]
            
        except Exception as e:
            raise Exception(f"Erro ao listar petições: {str(e)}")
    
    def get_petition_file(self, petition_id):
        """Retorna o caminho do arquivo de uma petição para download"""
        try:
            petition = GeneratedPetition.query.get(petition_id)
            if not petition:
                raise Exception("Petição não encontrada")
            
            return self.download_file(petition.gcs_path)
            
        except Exception as e:
            raise Exception(f"Erro ao obter arquivo da petição: {str(e)}")

document_service = DocumentService()